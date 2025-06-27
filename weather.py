"""
个性化职业发展与城市选择“探路者” (Career Pathfinder)


功能:
1.  图形用户界面 (GUI)。
2.  稳定的本地数据源 (`jobs.csv`)。
3.  多城市市场对比分析 (雷达图)。
4.  基于Dijkstra算法的技能路径规划 (网络图)。
5.  一键生成包含所有图文信息的Word分析报告。
6.  调用本地DeepSeek大模型，在报告中自动生成AI职业规划建议。
"""

# --- 核心库导入 ---
import tkinter as tk
from tkinter import ttk, messagebox
import json
import threading
import queue
import io
from datetime import datetime

# --- 第三方库导入 ---
import pandas as pd
import numpy as np
import networkx as nx
from scipy.sparse import csr_matrix
from scipy.sparse.csgraph import dijkstra
from docx import Document
from docx.shared import Inches
import requests

import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

# Matplotlib 全局中文字体设置
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False


# ==============================================================================
# 模块一: 后端数据处理与分析 (Backend Logic)
# ==============================================================================
class Backend:
    def _load_city_jobs_from_csv(self, job_title, city, result_queue):
        print(f"本地模式: 正在从 jobs.csv 加载 {city} 的 {job_title} 职位...")
        job_list = []
        try:
            df = pd.read_csv("jobs.csv")
            city_df = df[df['location'].str.contains(city, case=False, na=False)]
            keyword_df = city_df[city_df['title'].str.contains(job_title, case=False, na=False)]
            job_list = keyword_df.to_dict('records')
        except FileNotFoundError:
            print("错误: 未在项目目录中找到 jobs.csv 文件。请创建该文件。")
            job_list = []
        except Exception as e:
            print(f"读取 jobs.csv 时发生错误: {e}")
            job_list = []
        result_queue.put({city: job_list})
        print(f"本地模式: {city} 加载完成，共 {len(job_list)} 条数据。")

    def run_multi_city_load(self, job_title, cities):
        threads, result_queue = [], queue.Queue()
        for city in cities:
            thread = threading.Thread(target=self._load_city_jobs_from_csv, args=(job_title, city, result_queue))
            threads.append(thread)
            thread.start()
        for thread in threads:
            thread.join()
        all_results = {}
        while not result_queue.empty():
            all_results.update(result_queue.get())
        return all_results

    def analyze_market_data(self, all_jobs_data, user_skills):
        analysis_results, all_job_tags = {}, []
        for city, jobs in all_jobs_data.items():
            if not jobs: continue
            df = pd.DataFrame(jobs)
            job_count = len(df)
            skill_match_count = df['tags'].apply(
                lambda tags: sum(1 for skill in user_skills if str(skill).lower() in str(tags).lower())).sum()
            skill_match_ratio = skill_match_count / (job_count * len(user_skills)) if job_count and user_skills else 0
            analysis_results[city] = {'job_count': job_count, 'skill_match_ratio': skill_match_ratio * 100}
            all_job_tags.extend(df['tags'].tolist())
        return analysis_results, all_job_tags

    def build_and_analyze_skill_graph(self, all_job_tags, user_skills, target_skill):
        if not all_job_tags or not target_skill: return None, None
        all_skills_list = [skill.strip().lower() for tg in all_job_tags for skill in str(tg).split() if
                           len(skill.strip()) > 1]
        skill_counts = pd.Series(all_skills_list).value_counts()
        top_skills = skill_counts[skill_counts > 0].index.tolist()
        if not top_skills: return None, None
        all_nodes = list(set(top_skills + [s.lower() for s in user_skills] + [target_skill.lower()]))
        skill_to_idx = {skill: i for i, skill in enumerate(all_nodes)}
        co_occurrence = np.zeros((len(all_nodes), len(all_nodes)))
        for tag_group in all_job_tags:
            group_skills = list(
                set([s.strip().lower() for s in str(tag_group).split() if s.strip().lower() in skill_to_idx]))
            for i in range(len(group_skills)):
                for j in range(i + 1, len(group_skills)):
                    u, v = skill_to_idx[group_skills[i]], skill_to_idx[group_skills[j]]
                    co_occurrence[u, v] += 1
                    co_occurrence[v, u] += 1
        weights = np.where(co_occurrence > 0, 1 / (1 + co_occurrence), 0)
        graph = csr_matrix(weights)
        start_indices = [skill_to_idx.get(s.lower()) for s in user_skills if s.lower() in skill_to_idx]
        target_idx = skill_to_idx.get(target_skill.lower())
        if not start_indices or target_idx is None: return all_nodes, co_occurrence, None
        dist_matrix, predecessors = dijkstra(csgraph=graph, directed=False, indices=start_indices,
                                             return_predecessors=True)
        min_dist, path_start_node_idx, best_predecessors = np.inf, -1, None
        if dist_matrix.ndim > 1:
            if dist_matrix.shape[1] > target_idx and np.any(np.isfinite(dist_matrix[:, target_idx])):
                best_start_idx_in_list = np.argmin(dist_matrix[:, target_idx])
                min_dist = dist_matrix[best_start_idx_in_list, target_idx]
                path_start_node_idx = start_indices[best_start_idx_in_list]
                best_predecessors = predecessors[best_start_idx_in_list]
        else:
            if len(dist_matrix) > target_idx and np.isfinite(dist_matrix[target_idx]):
                min_dist, path_start_node_idx, best_predecessors = dist_matrix[target_idx], start_indices[
                    0], predecessors
        if np.isinf(min_dist): return all_nodes, co_occurrence, []
        path, curr = [], target_idx
        while curr != -9999 and curr != path_start_node_idx:
            path.append(all_nodes[curr])
            if best_predecessors is None or curr < 0 or curr >= len(best_predecessors): break
            curr = best_predecessors[curr]
        path.append(all_nodes[path_start_node_idx])
        return all_nodes, co_occurrence, path[::-1]

    def _get_ai_career_advice(self, prompt):
        print("正在向本地AI模型请求职业建议...")
        try:
            url = "http://localhost:11434/api/generate"
            payload = {"model": "deepseek-coder", "prompt": prompt, "stream": False}
            response = requests.post(url, json=payload, timeout=120)
            response.raise_for_status()
            response_data = response.json()
            return response_data.get("response", "AI模型未能返回有效建议。")
        except requests.exceptions.ConnectionError:
            return "错误：无法连接到本地Ollama服务。请确保Ollama正在运行，并且您已通过 `ollama run deepseek-coder` 加载了模型。"
        except Exception as e:
            return f"与AI模型通信时发生错误: {e}"

    def generate_word_report(self, profile_info, analysis_results, skill_path_info, fig_dashboard, fig_skill_net):
        try:
            doc = Document()
            doc.add_heading('个性化职业发展与城市选择分析报告', level=0)
            doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_heading('一、您的个人档案', level=1)
            doc.add_paragraph(f"    目标职位关键字: {profile_info['job_title']}")
            doc.add_paragraph(f"    核心技能: {profile_info['skills']}")
            doc.add_paragraph(f"    目标学习技能: {profile_info['target_skill']}")
            doc.add_heading('二、多城市市场分析总结', level=1)
            summary_text = ""
            if not analysis_results:
                summary_text = "未找到相关数据，无法生成分析。\n"
                doc.add_paragraph(summary_text)
            else:
                for city, result in analysis_results.items():
                    line = (f"在城市 {city}，共找到 {result['job_count']} 个相关职位。"
                            f"您的核心技能与该城市职位的匹配度约为 {result['skill_match_ratio']:.2f}%。\n")
                    summary_text += line
                    doc.add_paragraph(line, style='List Bullet')
            doc.add_heading('三、可视化图表', level=1)
            doc.add_paragraph("下图展示了所选城市的综合竞争力对比：")
            memfile_dashboard = io.BytesIO()
            fig_dashboard.savefig(memfile_dashboard, format='png', dpi=300)
            memfile_dashboard.seek(0)
            doc.add_picture(memfile_dashboard, width=Inches(6.0))
            doc.add_paragraph("\n下图展示了相关技能的共现网络，以及为您推荐的学习路径（红色高亮）：")
            memfile_skill_net = io.BytesIO()
            fig_skill_net.savefig(memfile_skill_net, format='png', dpi=300)
            memfile_skill_net.seek(0)
            doc.add_picture(memfile_skill_net, width=Inches(6.0))
            doc.add_heading('四、AI 智能职业规划建议', level=1)
            prompt = (
                f"你是一位资深的职业规划专家。请根据以下我的个人情况和市场分析数据，为我生成一段150-200字的、个性化的职业发展建议。\n\n"
                f"--- 我的档案 ---\n目标职位: {profile_info['job_title']}\n我的核心技能: {profile_info['skills']}\n我希望学习的下一个技能: {profile_info['target_skill']}\n\n"
                f"--- 市场数据摘要 ---\n{summary_text}\n"
                f"--- 推荐学习路径 ---\n数据分析得出的、从我现有技能到目标技能的推荐路径是: {' -> '.join(skill_path_info)}\n\n"
                f"--- 你的任务 ---\n请结合以上所有信息，用亲切、专业的口吻，给我提供一段具体的行动建议，包括：\n1. 基于数据，简要评价我选择的城市是否合适。\n2. 解释为什么推荐的学习路径是合理的。\n3. 给我一个明确的、可以立即开始的下一步行动建议。")
            ai_advice = self._get_ai_career_advice(prompt)
            doc.add_paragraph(ai_advice)
            file_name = f"AI智能职业分析报告_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(file_name)
            return f"报告已成功生成！已保存为 {file_name}"
        except Exception as e:
            return f"生成报告时发生错误: {e}"


# ==============================================================================
# 模块二: 前端GUI与可视化 (Frontend & Visualization)
# ==============================================================================
class PathfinderApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("职业发展与城市选择“探路者” V3.0 (AI增强版)")
        self.geometry("1280x800")
        self.configure(bg='#F0F0F0')

        self.backend = Backend()
        self.analysis_results, self.skill_graph_data = None, None

        self._define_fonts()
        self._configure_styles()
        self._create_widgets()
        self._load_profile()

    def _define_fonts(self):
        self.font_title = ("Microsoft YaHei", 12, "bold")
        self.font_label = ("Microsoft YaHei", 10)
        self.font_status = ("Microsoft YaHei", 9, "italic")

    def _configure_styles(self):
        style = ttk.Style(self)
        style.theme_use('clam')
        style.configure('TFrame', background='#F0F0F0')
        style.configure('TLabel', background='#F0F0F0', font=self.font_label)
        style.configure('TButton', font=self.font_label, padding=5)
        style.configure('TEntry', padding=5)
        style.configure('TCheckbutton', background='#F0F0F0', font=self.font_label)
        style.configure('TLabelframe', background='#F0F0F0', padding=10)
        style.configure('TLabelframe.Label', background='#F0F0F0', font=self.font_title, foreground='#003366')
        style.configure('Primary.TButton', background='#0078D7', foreground='white')
        style.map('Primary.TButton', background=[('active', '#005a9e')])

    def _create_widgets(self):
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        control_panel = ttk.Frame(main_frame, width=320, style='TFrame')
        control_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        control_panel.pack_propagate(False)
        display_panel = ttk.Frame(main_frame)
        display_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        self._create_profile_frame(control_panel)
        self._create_market_frame(control_panel)
        self._create_skill_path_frame(control_panel)
        self._create_status_bar(control_panel)
        notebook = ttk.Notebook(display_panel)
        notebook.pack(fill=tk.BOTH, expand=True)
        self.tab1, self.tab2 = ttk.Frame(notebook), ttk.Frame(notebook)
        notebook.add(self.tab1, text='城市对比仪表盘')
        notebook.add(self.tab2, text='技能网络与路径')
        self.fig = Figure(figsize=(10, 8), dpi=100, facecolor='#FFFFFF')
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.tab1)
        self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        self.fig2 = Figure(figsize=(10, 8), dpi=100, facecolor='#FFFFFF')
        self.canvas2 = FigureCanvasTkAgg(self.fig2, master=self.tab2)
        self.canvas2.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def _create_profile_frame(self, parent):
        frame = ttk.LabelFrame(parent, text="我的档案")
        frame.pack(fill=tk.X, padx=5, pady=10)
        ttk.Label(frame, text="目标职位:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.job_title_entry = ttk.Entry(frame, width=25, font=self.font_label)
        self.job_title_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        ttk.Label(frame, text="核心技能:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.skills_entry = ttk.Entry(frame, width=25, font=self.font_label)
        self.skills_entry.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        save_btn = ttk.Button(frame, text="保存档案", command=self._save_profile)
        save_btn.grid(row=2, column=0, columnspan=2, pady=10)
        frame.columnconfigure(1, weight=1)

    def _create_market_frame(self, parent):
        frame = ttk.LabelFrame(parent, text="市场分析")
        frame.pack(fill=tk.X, padx=5, pady=10)
        ttk.Label(frame, text="选择对比城市:").pack(anchor="w", padx=5)
        cities_frame = ttk.Frame(frame)
        cities_frame.pack(fill=tk.X, padx=15)
        default_cities = ["New York", "London", "San Francisco", "Berlin", "Sydney", "Toronto"]
        self.city_vars = {}
        for i, city in enumerate(default_cities):
            var = tk.BooleanVar(value=(city in ["New York", "London", "San Francisco"]))
            self.city_vars[city] = var
            ttk.Checkbutton(cities_frame, text=city, variable=var).grid(row=i // 2, column=i % 2, sticky='w')
        self.analyze_btn = ttk.Button(frame, text="开始分析", command=self.run_analysis, style='Primary.TButton')
        self.analyze_btn.pack(pady=10, fill=tk.X, padx=5)
        self.report_btn = ttk.Button(frame, text="生成AI分析报告", command=self.generate_report, state=tk.DISABLED)
        self.report_btn.pack(pady=5, fill=tk.X, padx=5)

    def _create_skill_path_frame(self, parent):
        frame = ttk.LabelFrame(parent, text="技能路径规划")
        frame.pack(fill=tk.X, padx=5, pady=10)
        ttk.Label(frame, text="目标技能:").pack(anchor="w", padx=5)
        self.target_skill_entry = ttk.Entry(frame, width=25, font=self.font_label)
        self.target_skill_entry.pack(pady=5, padx=5, fill=tk.X)
        self.plan_path_btn = ttk.Button(frame, text="规划学习路径", command=self.plan_skill_path, state=tk.DISABLED)
        self.plan_path_btn.pack(pady=10, fill=tk.X, padx=5)

    def _create_status_bar(self, parent):
        status_frame = ttk.Frame(parent, height=30)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=5)
        self.status_label = ttk.Label(status_frame, text="准备就绪。", font=self.font_status, relief=tk.SUNKEN,
                                      anchor=tk.W, padding=5)
        self.status_label.pack(fill=tk.BOTH, expand=True)

    def _save_profile(self):
        profile = {"job_title": self.job_title_entry.get(), "skills": self.skills_entry.get(),
                   "target_skill": self.target_skill_entry.get()}
        with open("profile.json", "w", encoding="utf-8") as f: json.dump(profile, f, ensure_ascii=False, indent=4)
        self.status_label.config(text="档案已保存到 profile.json")

    def _load_profile(self):
        try:
            with open("profile.json", "r", encoding="utf-8") as f:
                profile = json.load(f)
                self.job_title_entry.insert(0, profile.get("job_title", "Python Developer"))
                self.skills_entry.insert(0, profile.get("skills", "python,pandas,sql"))
                self.target_skill_entry.insert(0, profile.get("target_skill", "scikit-learn"))
                self.status_label.config(text="已成功加载档案。")
        except FileNotFoundError:
            self.status_label.config(text="未找到档案文件，请输入您的信息。")

    def run_analysis(self):
        job_title = self.job_title_entry.get()
        selected_cities = [city for city, var in self.city_vars.items() if var.get()]
        if not job_title or not selected_cities:
            messagebox.showerror("输入错误", "请输入目标职位并至少选择一个城市。")
            return
        self.analyze_btn.config(state=tk.DISABLED);
        self.plan_path_btn.config(state=tk.DISABLED);
        self.report_btn.config(state=tk.DISABLED)
        self.status_label.config(text="正在分析，请稍候...")
        threading.Thread(target=self._analysis_worker, args=(job_title, selected_cities), daemon=True).start()

    def _analysis_worker(self, job_title, cities):
        self.status_label.config(text=f"正在从本地文件加载 {len(cities)} 个城市的数据...")
        all_jobs = self.backend.run_multi_city_load(job_title, cities)
        self.status_label.config(text="正在分析市场数据...")
        user_skills = [s.strip() for s in self.skills_entry.get().split(',') if s.strip()]
        self.analysis_results, all_job_tags = self.backend.analyze_market_data(all_jobs, user_skills)
        self.status_label.config(text="正在构建技能关系图...")
        target_skill = self.target_skill_entry.get()
        self.skill_graph_data = self.backend.build_and_analyze_skill_graph(all_job_tags, user_skills, target_skill)
        self.after(0, self.update_dashboard)
        self.after(0, self.plan_skill_path)

    def update_dashboard(self):
        if not self.analysis_results or not any(self.analysis_results.values()):
            self.status_label.config(text="分析完成，但未找到有效数据。");
            self.analyze_btn.config(state=tk.NORMAL)
            return
        self.fig.clear();
        ax = self.fig.add_subplot(111, polar=True)
        cities = list(self.analysis_results.keys());
        labels = ['职位数量', '技能匹配度', '市场热度']
        max_vals = {
            'job_count': max((d['job_count'] for d in self.analysis_results.values() if d and d.get('job_count')),
                             default=1),
            'skill_match_ratio': max(
                (d['skill_match_ratio'] for d in self.analysis_results.values() if d and d.get('skill_match_ratio')),
                default=1)}
        data, valid_cities = [], []
        for city in cities:
            res = self.analysis_results.get(city)
            if not res: continue
            valid_cities.append(city)
            job_count_norm = res['job_count'] / max_vals['job_count']
            skill_match_norm = res['skill_match_ratio'] / max_vals['skill_match_ratio']
            market_heat_norm = np.sqrt(job_count_norm)
            data.append([job_count_norm, skill_match_norm, market_heat_norm])
        angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist();
        angles += angles[:1]
        ax.set_thetagrids(np.degrees(angles[:-1]), labels)
        for i, city_data in enumerate(data):
            city_data += city_data[:1];
            ax.plot(angles, city_data, label=valid_cities[i], marker='o');
            ax.fill(angles, city_data, alpha=0.25)
        ax.set_title("城市职位市场竞争力雷达图");
        ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.1))
        self.fig.tight_layout();
        self.canvas.draw()
        self.status_label.config(text="分析完成！");
        self.analyze_btn.config(state=tk.NORMAL);
        self.plan_path_btn.config(state=tk.NORMAL);
        self.report_btn.config(state=tk.NORMAL)

    def plan_skill_path(self):
        if self.skill_graph_data is None or len(self.skill_graph_data) != 3:
            messagebox.showinfo("提示", "未能构建有效的技能网络图。");
            return
        nodes, matrix, path = self.skill_graph_data;
        self.fig2.clear();
        ax = self.fig2.add_subplot(111)
        if nodes is None:
            ax.text(0.5, 0.5, "未能构建有效的技能网络图。", ha='center', va='center');
            self.canvas2.draw();
            return
        G = nx.from_numpy_array(matrix);
        mapping = {i: node for i, node in enumerate(nodes)};
        G = nx.relabel_nodes(G, mapping)
        pos = nx.spring_layout(G, k=0.8, iterations=50)
        nx.draw_networkx_nodes(G, pos, node_size=50, node_color='skyblue', ax=ax);
        nx.draw_networkx_edges(G, pos, alpha=0.3, ax=ax)
        if path:
            path_edges = list(zip(path, path[1:]))
            nx.draw_networkx_nodes(G, pos, nodelist=path, node_color='tomato', node_size=150, ax=ax)
            nx.draw_networkx_edges(G, pos, edgelist=path_edges, edge_color='red', width=2, ax=ax)
        high_degree_nodes = {node for node, degree in G.degree() if degree > 2 or node in (path or [])}
        labels = {node: node for node in high_degree_nodes}
        nx.draw_networkx_labels(G, pos, labels=labels, font_size=8, ax=ax)
        ax.set_title("技能网络图与推荐学习路径 (红色)");
        self.canvas2.draw()

    def generate_report(self):
        if not self.analysis_results: messagebox.showwarning("操作无效",
                                                             "请先执行一次成功的市场分析，再生成报告。"); return
        if self.skill_graph_data is None or len(self.skill_graph_data) != 3: messagebox.showwarning("数据不完整",
                                                                                                    "技能路径数据尚未生成，请先规划一次学习路径。"); return
        self.report_btn.config(state=tk.DISABLED);
        self.status_label.config(text="正在生成AI智能报告，请稍候...")
        profile_info = {"job_title": self.job_title_entry.get(), "skills": self.skills_entry.get(),
                        "target_skill": self.target_skill_entry.get()}
        _, _, skill_path = self.skill_graph_data
        threading.Thread(target=self._report_worker,
                         args=(profile_info, self.analysis_results, skill_path, self.fig, self.fig2),
                         daemon=True).start()

    def _report_worker(self, profile_info, analysis_results, skill_path, fig1, fig2):
        message = self.backend.generate_word_report(profile_info, analysis_results, skill_path, fig1, fig2)
        self.after(0, self.on_report_generated, message)

    def on_report_generated(self, message):
        self.status_label.config(text="报告生成完毕！");
        self.report_btn.config(state=tk.NORMAL)
        messagebox.showinfo("操作成功", message)


if __name__ == '__main__':
    app = PathfinderApp()
    app.mainloop()